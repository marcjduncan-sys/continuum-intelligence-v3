"""
Text extraction from uploaded research documents.

Supports PDF (via PyMuPDF), DOCX (via python-docx), and plain text files.
All extraction is in-memory; no files are written to disk.
"""

import io
import os
import re

# Reuse the HTML entity cleanup from the ingest pipeline
from ingest import _clean_html


def extract_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a PDF file.

    Returns (extracted_text, page_count).
    Raises ValueError if the PDF contains only scanned images (no text).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        page_count = len(doc)
        pages_text = []
        for page in doc:
            text = page.get_text()
            pages_text.append(text)

        if not any(t.strip() for t in pages_text):
            raise ValueError(
                "This PDF appears to contain scanned images rather than text. "
                "Text-based PDFs and DOCX files are supported."
            )

        combined = "\n\n".join(pages_text)
    finally:
        doc.close()

    combined = _clean_html(combined)
    # Collapse 3+ newlines to 2
    combined = re.sub(r"\n{3,}", "\n\n", combined)
    return combined, page_count


def extract_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a DOCX file.

    Returns (extracted_text, page_count_estimate).
    Page count is estimated as total_chars / 2500, minimum 1.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = "\t".join(cells)
            if row_text.strip():
                parts.append(row_text)

    combined = "\n\n".join(parts)
    combined = _clean_html(combined)
    combined = re.sub(r"\n{3,}", "\n\n", combined)

    char_count = len(combined)
    page_estimate = max(1, char_count // 2500)

    return combined, page_estimate


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, int, str]:
    """Dispatch to the correct extractor based on file extension.

    Returns (extracted_text, page_count, mime_type).
    Raises ValueError for unsupported file types.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text, pages = extract_pdf(file_bytes)
        return text, pages, "application/pdf"

    if ext == ".docx":
        text, pages = extract_docx(file_bytes)
        return (
            text,
            pages,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if ext in (".txt", ".md"):
        text = file_bytes.decode("utf-8", errors="replace")
        text = _clean_html(text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        page_estimate = max(1, len(text) // 2500)
        return text, page_estimate, "text/plain"

    raise ValueError(
        f"Unsupported file type: '{ext}'. Supported formats: PDF, DOCX, TXT."
    )


_MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def validate_upload(file_bytes: bytes, filename: str) -> None:
    """Pre-extraction validation.

    Raises ValueError if file is empty, too large, or unsupported type.
    """
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")

    if len(file_bytes) > _MAX_FILE_SIZE:
        size_mb = len(file_bytes) / (1024 * 1024)
        raise ValueError(
            f"File size ({size_mb:.1f} MB) exceeds the 10 MB limit."
        )

    ext = os.path.splitext(filename)[1].lower()
    if ext not in _ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: '{ext}'. Supported formats: PDF, DOCX, TXT."
        )
