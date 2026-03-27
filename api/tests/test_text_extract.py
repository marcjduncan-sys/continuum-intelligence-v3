"""Tests for text extraction from PDF, DOCX, and plain text files."""

import io
import sys
import os

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from text_extract import (
    extract_docx,
    extract_pdf,
    extract_text,
    validate_upload,
)


# ---------------------------------------------------------------------------
# Helpers: generate minimal test fixtures programmatically
# ---------------------------------------------------------------------------

def _make_pdf_with_text(text: str = "Hello from a test PDF.") -> bytes:
    """Create a minimal single-page PDF with text using PyMuPDF."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_pdf_image_only() -> bytes:
    """Create a PDF with a coloured rectangle but no text."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    # Draw a filled rectangle (no text)
    rect = fitz.Rect(50, 50, 200, 200)
    page.draw_rect(rect, color=(0, 0, 0), fill=(0.5, 0.5, 0.5))
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def _make_docx_with_paragraphs(paragraphs: list[str]) -> bytes:
    """Create a minimal DOCX with the given paragraphs."""
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(rows: list[list[str]]) -> bytes:
    """Create a DOCX containing a table."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Table below:")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.cell(i, j).text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 1. extract_pdf returns text and page count from a valid PDF
# ---------------------------------------------------------------------------

def test_extract_pdf_valid():
    pdf_bytes = _make_pdf_with_text("Revenue grew 12% year over year.")
    text, pages = extract_pdf(pdf_bytes)
    assert "Revenue grew 12%" in text
    assert pages == 1


# ---------------------------------------------------------------------------
# 2. extract_pdf raises ValueError for image-only PDF
# ---------------------------------------------------------------------------

def test_extract_pdf_image_only():
    pdf_bytes = _make_pdf_image_only()
    with pytest.raises(ValueError, match="scanned images"):
        extract_pdf(pdf_bytes)


# ---------------------------------------------------------------------------
# 3. extract_docx returns text from paragraphs
# ---------------------------------------------------------------------------

def test_extract_docx_paragraphs():
    docx_bytes = _make_docx_with_paragraphs([
        "First paragraph about earnings.",
        "Second paragraph about risks.",
    ])
    text, pages = extract_docx(docx_bytes)
    assert "First paragraph about earnings." in text
    assert "Second paragraph about risks." in text
    assert pages >= 1


# ---------------------------------------------------------------------------
# 4. extract_docx returns text from tables
# ---------------------------------------------------------------------------

def test_extract_docx_tables():
    docx_bytes = _make_docx_with_table([
        ["Metric", "Value"],
        ["Revenue", "$100M"],
        ["EBITDA", "$25M"],
    ])
    text, pages = extract_docx(docx_bytes)
    assert "Revenue" in text
    assert "$100M" in text
    assert "EBITDA" in text


# ---------------------------------------------------------------------------
# 5. extract_text dispatches correctly by extension
# ---------------------------------------------------------------------------

def test_extract_text_pdf():
    pdf_bytes = _make_pdf_with_text("Test dispatch.")
    text, pages, mime = extract_text(pdf_bytes, "report.pdf")
    assert "Test dispatch" in text
    assert mime == "application/pdf"


def test_extract_text_txt():
    content = b"Plain text research note content."
    text, pages, mime = extract_text(content, "notes.txt")
    assert "Plain text research" in text
    assert mime == "text/plain"


def test_extract_text_md():
    content = b"# Markdown heading\n\nSome analysis."
    text, pages, mime = extract_text(content, "analysis.md")
    assert "Markdown heading" in text
    assert mime == "text/plain"


# ---------------------------------------------------------------------------
# 6. extract_text raises ValueError for unsupported extension
# ---------------------------------------------------------------------------

def test_extract_text_unsupported():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(b"data", "spreadsheet.xlsx")


# ---------------------------------------------------------------------------
# 7. validate_upload passes for valid PDF under 10MB
# ---------------------------------------------------------------------------

def test_validate_upload_valid_pdf():
    pdf_bytes = _make_pdf_with_text("Valid.")
    validate_upload(pdf_bytes, "report.pdf")  # Should not raise


def test_validate_upload_valid_docx():
    docx_bytes = _make_docx_with_paragraphs(["Valid."])
    validate_upload(docx_bytes, "report.docx")  # Should not raise


# ---------------------------------------------------------------------------
# 8. validate_upload raises for file > 10MB
# ---------------------------------------------------------------------------

def test_validate_upload_too_large():
    large_bytes = b"x" * (10 * 1024 * 1024 + 1)
    with pytest.raises(ValueError, match="10 MB limit"):
        validate_upload(large_bytes, "huge.pdf")


# ---------------------------------------------------------------------------
# 9. validate_upload raises for empty file
# ---------------------------------------------------------------------------

def test_validate_upload_empty():
    with pytest.raises(ValueError, match="empty"):
        validate_upload(b"", "empty.pdf")


# ---------------------------------------------------------------------------
# 10. validate_upload raises for unsupported extension
# ---------------------------------------------------------------------------

def test_validate_upload_unsupported_ext():
    with pytest.raises(ValueError, match="Unsupported file type"):
        validate_upload(b"data", "file.xlsx")

    with pytest.raises(ValueError, match="Unsupported file type"):
        validate_upload(b"data", "slides.pptx")
