#!/usr/bin/env python3
"""
Branded Investment Report DOCX Generator — Continuum Trinity
Generates visually branded DOCX reports replicating the HTML investment report design.

Usage:
    python scripts/generate_investment_report_docx.py output.docx --format short
    python scripts/generate_investment_report_docx.py output.docx --format long
"""

import argparse
import sys
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════════════
# BRAND COLOURS
# ═══════════════════════════════════════════════════════════════

C = {
    "deep_teal":      RGBColor(0x1A, 0x5F, 0x6C),
    "midnight":       RGBColor(0x0E, 0x3A, 0x42),
    "sage":           RGBColor(0x4A, 0x9E, 0x7E),
    "gold":           RGBColor(0xC9, 0xA9, 0x6E),
    "red":            RGBColor(0xC5, 0x30, 0x30),
    "amber":          RGBColor(0xC0, 0x7A, 0x1A),
    "green":          RGBColor(0x2F, 0x85, 0x5A),
    "navy":           RGBColor(0x1B, 0x2A, 0x4A),
    "charcoal":       RGBColor(0x2D, 0x34, 0x40),
    "text_primary":   RGBColor(0x1A, 0x20, 0x2C),
    "text_secondary": RGBColor(0x4A, 0x55, 0x68),
    "text_muted":     RGBColor(0x88, 0x98, 0xAA),
    "white":          RGBColor(0xFF, 0xFF, 0xFF),
    "border":         RGBColor(0xE2, 0xE8, 0xF0),
    "light_green":    RGBColor(0x68, 0xD3, 0x91),
    "red_light":      RGBColor(0xFC, 0x81, 0x81),
}

HEX = {
    "midnight":   "0E3A42",
    "deep_teal":  "1A5F6C",
    "sidebar":    "F3F5F7",
    "alt_row":    "F7F9FA",
    "callout":    "F0F7F8",
    "callout_warn": "FFF8F0",
    "callout_red":  "FFF0F0",
    "callout_green": "F0FFF4",
    "border":     "E2E8F0",
    "table_hdr":  "1A5F6C",
    "white":      "FFFFFF",
}

HYP_COLOURS = {
    "H1": (C["green"],  "2F855A"),
    "H2": (C["amber"],  "C07A1A"),
    "H3": (C["red"],    "C53030"),
    "H4": (C["text_muted"], "8898AA"),
}

EP_BADGES = {
    "Motivated":     ("FFF3E0", RGBColor(0xE6, 0x51, 0x00)),
    "Under Oath":    ("E8F5E9", RGBColor(0x2E, 0x7D, 0x32)),
    "Under Oath / Statutory": ("E8F5E9", RGBColor(0x2E, 0x7D, 0x32)),
    "Consensus":     ("E3F2FD", RGBColor(0x15, 0x65, 0xC0)),
    "Independent":   ("F3E5F5", RGBColor(0x7B, 0x1F, 0xA2)),
    "Objective":     ("E0F2F1", RGBColor(0x00, 0x69, 0x5C)),
    "Behavioural":   ("FFF8E1", RGBColor(0xF5, 0x7F, 0x17)),
    "Peer-Reviewed": ("EDE7F6", RGBColor(0x45, 0x27, 0xA0)),
    "Noise":         ("EFEBE9", RGBColor(0x5D, 0x40, 0x37)),
}

COV_BADGES = {
    "Full":    ("E8F5E9", RGBColor(0x2E, 0x7D, 0x32)),
    "Good":    ("E0F2F1", RGBColor(0x4A, 0x9E, 0x7E)),
    "Partial": ("FFF3E0", RGBColor(0xC0, 0x7A, 0x1A)),
    "Limited": ("F5F5F5", RGBColor(0x88, 0x98, 0xAA)),
}

FONT = "Calibri"
MONO = "Consolas"


# ═══════════════════════════════════════════════════════════════
# OOXML UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, fill_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_width(cell, width):
    cell.width = width

def set_cell_margins(cell, top=0, bottom=0, left=72, right=72):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    val = int(height_pt * 20)
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{val}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def set_para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill_hex}"/>')
    pPr.append(shd)

def set_para_border_left(para, colour_hex, sz=18, space=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="{space}" w:color="{colour_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def set_para_border_bottom(para, colour_hex, sz=6, space=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="{space}" w:color="{colour_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def set_run_shading(run, fill_hex):
    rPr = run._r.get_or_add_rPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill_hex}"/>')
    rPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_borders(table, colour_hex="E2E8F0", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{colour_hex}"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    val_map = {"center": "center", "top": "top", "bottom": "bottom"}
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val_map.get(align, "center")}"/>')
    tcPr.append(vAlign)


def add_run(para, text, size=10, colour=None, bold=False, italic=False, font=None, strike=False):
    run = para.add_run(text)
    run.font.name = font or FONT
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    run.font.bold = bold
    run.font.italic = italic
    if strike:
        run.font.strike = True
    return run

def add_badge_run(para, text, fill_hex, text_colour, size=7):
    run = para.add_run(f" {text} ")
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = text_colour
    set_run_shading(run, fill_hex)
    return run

def spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    pf = p.paragraph_format
    pf.line_spacing = Pt(pts)
    r = p.add_run("")
    r.font.size = Pt(1)
    return p

def section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text.upper(), size=7, colour=C["deep_teal"], bold=True)
    return p


# ═══════════════════════════════════════════════════════════════
# WOW.AX REPORT DATA
# ═══════════════════════════════════════════════════════════════

DATA = {
    "ticker": "WOW.AX",
    "company": "Woolworths Group Limited",
    "sector": "Consumer Staples",
    "subsector": "Supermarkets",
    "date": "10 February 2026",
    "version": "v1.0",
    "report_id": "WOW-2026-001",
    "mode": "Narrative Intelligence",
    "price": "A$31.41",
    "next_update": "March 2026 (or post 25-Feb results)",

    "risk_skew": "Downside",
    "risk_skew_rationale": (
        "Evidence accumulating against the embedded narrative. 5 of 8 domains support "
        "erosion (H2); price above consensus target. No valuation cushion if H1 FY26 "
        "results disappoint on 25 Feb."
    ),

    "metrics": [
        ("Mkt Cap", "A$38.3B", ""),
        ("Fwd P/E", "23.5x", "premium"),
        ("EV/EBITDA", "10.5x", ""),
        ("Div Yield", "2.9%", ""),
        ("NPAT (FY25)", "\u219319%", "negative"),
        ("EBIT Margin", "4.0% (\u219370bp)", "negative"),
        ("Mkt Share", "38% (Duopoly 67%)", ""),
        ("Next Event", "H1 FY26: 25 Feb", ""),
    ],

    "narrative": (
        "The market is pricing a turnaround under new CEO Amanda Bardwell. "
        "23.5x forward P/E assumes $400M cost-out program delivers, e-commerce scales "
        "(penetration now 16.2%), and Australian Food returns to mid-to-high single-digit "
        'EBIT growth. FY25 was "disappointing"; FY26 is framed as "transitional."'
    ),

    "verdict": (
        "Turnaround narrative is plausible but unproven. $95M industrial action cost was "
        "one-off, but structural headwinds (mandatory FGCC, ACCC court proceedings, margin "
        "pressure from price investment, Aldi expansion) are not cyclical. Market pricing "
        "recovery that management itself calls uncertain."
    ),

    "verdict_long": (
        "The market is pricing a Bardwell turnaround at a premium multiple, but FY25 "
        'delivered a 19% NPAT decline. FY26 is "transitional"; the H1 results on '
        "25 February are the single most diagnostic event for whether this is a trough "
        "or a new normal."
    ),

    "hypotheses": [
        {
            "id": "H1", "name": "Managed Turnaround", "score": 50,
            "direction": "Awaiting", "status": "Priced",
            "desc": "$400M cost-out delivers. E-commerce scales. Australian Food EBIT grows mid-high single digits. Dividend recovers.",
            "full_desc": (
                "$400M cost-out delivers margin recovery. E-commerce scales profitably. "
                "Australian Food EBIT growth returns to mid-high single digits. Big W returns to profit. "
                "New CEO resets operational culture. Dividend recovers toward A$1.00+."
            ),
            "requires": [
                "Australian Food EBIT growth of 8% or above in FY26 (consensus ~11.8%)",
                "$400M cost savings not fully consumed by price investment or wage costs",
                "E-commerce penetration exceeds 18% without margin dilution",
                "ACCC court outcome limited: settlement or manageable penalty",
                "Big W EBIT returns to breakeven or better",
            ],
            "supporting": [
                "Corporate Comms: Q1 FY26 trading positive: sales +2.1%, NPS +3, value perception +5",
                "Broker: Consensus Hold/Buy with expected dividend recovery",
                "Alternative Data: E-commerce infrastructure investment (Auburn CFC, Moorebank NDC) is real capex commitment",
                "Economic: RBA rate cuts improving consumer sentiment; food inflation moderating",
            ],
            "contradicting": [
                "Regulatory: Three concurrent regulatory actions create structural headwinds",
                "Media: Brand trust deterioration and political targeting",
                'Management: CEO herself says current performance "below our ambition"',
            ],
        },
        {
            "id": "H2", "name": "Structural Margin Erosion", "score": 35,
            "direction": "Rising", "status": "Evidence Building",
            "desc": "Price investment to defend share permanently compresses margins. Duopoly premium erodes. 4% EBIT margin becomes ceiling, not floor.",
            "full_desc": (
                "Price investment to defend market share against Aldi and cost-conscious consumers permanently "
                "compresses margins. The 4.0% EBIT margin in FY25 becomes a ceiling, not a trough. $400M cost savings "
                "are consumed by price investment, wage increases (11% over 3 years from strike settlement), and "
                "e-commerce fulfilment costs. The duopoly premium erodes as competition intensifies."
            ),
            "requires": [
                "Australian Food EBIT margin fails to recover above 4.5% despite cost-out program",
                "Grocery deflation (7 consecutive quarters of price cuts) continues",
                "Wage inflation runs ahead of productivity gains",
                "E-commerce fulfilment remains margin-dilutive at current penetration levels",
            ],
            "supporting": [
                "Regulatory: ACCC inquiry confirmed both chains increased product margins; now forced into price reversal",
                "Competitor: Aldi expansion forcing defensive pricing; duopoly dynamics weakening",
                "Economic: Wage growth above CPI; tobacco decline removing high-margin revenue",
                "Academic: Base rates for duopoly margin compression in grocery: 100-200bp over a decade",
                "Media: Disappointing results narrative; share price down 11.6% over 5 years",
                "Broker: Share price above consensus target: limited upside even in bull case",
            ],
            "contradicting": [
                "Corporate Comms: $400M cost-out positioned as enabling both price investment and margin recovery",
                "Alt Data: E-commerce infrastructure investment may unlock scale benefits at higher penetration",
            ],
        },
        {
            "id": "H3", "name": "Regulatory Squeeze", "score": 30,
            "direction": "Rising", "status": "Active",
            "desc": "Mandatory FGCC + ACCC court action + 20 inquiry recommendations compress supplier margins and impose compliance costs.",
            "full_desc": (
                "The cumulative regulatory burden (mandatory FGCC, ACCC court proceedings, 20 inquiry recommendations, "
                "class action, fresh produce provisions from April 2026) creates a structurally higher compliance cost base "
                "and constrains aggressive supplier negotiation that historically supported gross margins. Political salience "
                "of grocery pricing ensures sustained regulatory attention regardless of which party is in government."
            ),
            "requires": [
                "ACCC court outcome sets punitive precedent (penalties and/or structural orders)",
                "Mandatory FGCC compliance costs exceed $30M per annum",
                "Supplier renegotiation under new code constraints reduces gross margin by 20-50bp",
                "Political environment maintains pressure on supermarket pricing",
            ],
            "supporting": [
                "Regulatory: Mandatory code in force. ACCC proceedings commenced. 20 recommendations. Penalties: up to 10% of turnover",
                'Media: "Price gouging" narrative politically salient. Cost-of-living framing ensures bipartisan enforcement',
                "Economic: Consumer pressure increases political appetite for intervention",
            ],
            "contradicting": [
                "Academic: UK Groceries Code Adjudicator precedent shows moderate impact on aggregate profitability",
                "Regulatory: Emerson Review rejected calls for forced break-up; remedies may be limited to behavioural constraints",
            ],
        },
        {
            "id": "H4", "name": "Competitive Disruption", "score": 25,
            "direction": "Steady", "status": "Slow Burn",
            "desc": "Aldi 10%+ share and expanding. Amazon enters grocery. Online-only disruptors. Duopoly premium structurally threatened.",
            "full_desc": (
                "Aldi's physical expansion continues eroding duopoly share. Amazon enters grocery more aggressively. "
                "If Aldi launches online grocery (as it has in other markets), Woolworths' key competitive moat "
                "(e-commerce leadership) evaporates. Duopoly pricing power is structurally broken over a 5-10 year horizon."
            ),
            "requires": [],
            "supporting": [
                "Competitor: Aldi at 602 stores, growing. Targeting premium segments. Amazon pushing categories",
                "Academic: Discount entrant dynamics consistently erode incumbent share over decade-plus timeframes",
                "Alt Data: Industrial action around AI surveillance shows automation/workforce tension that could recur",
            ],
            "contradicting": [
                "Aldi has no online offering and no announced plans; Woolworths' e-com moat holds for now",
                "38% share and 1,000+ store network provides distribution advantage new entrants cannot replicate quickly",
                "Costco/Amazon are category-selective, not full-service competitors",
            ],
        },
    ],

    # Evidence matrix: domain -> {epistemic, signals: {H1, H2, H3, H4}}
    # Signals: "strong_support", "support", "weak_support", "contradict", "neutral"
    "evidence_matrix": [
        {"domain": "Corporate Comms", "epistemic": "Motivated",
         "signals": {"H1": "support", "H2": "neutral", "H3": "neutral", "H4": "neutral"}},
        {"domain": "Regulatory Filings", "epistemic": "Under Oath",
         "signals": {"H1": "contradict", "H2": "support", "H3": "strong_support", "H4": "neutral"}},
        {"domain": "Broker Research", "epistemic": "Consensus",
         "signals": {"H1": "support", "H2": "weak_support", "H3": "neutral", "H4": "neutral"}},
        {"domain": "Competitor Data", "epistemic": "Independent",
         "signals": {"H1": "neutral", "H2": "support", "H3": "neutral", "H4": "strong_support"}},
        {"domain": "Economic Data", "epistemic": "Objective",
         "signals": {"H1": "weak_support", "H2": "support", "H3": "support", "H4": "neutral"}},
        {"domain": "Alternative Data", "epistemic": "Behavioural",
         "signals": {"H1": "support", "H2": "neutral", "H3": "neutral", "H4": "support"}},
        {"domain": "Academic Research", "epistemic": "Peer-Reviewed",
         "signals": {"H1": "neutral", "H2": "support", "H3": "neutral", "H4": "support"}},
        {"domain": "Media & Social", "epistemic": "Noise",
         "signals": {"H1": "contradict", "H2": "support", "H3": "strong_support", "H4": "neutral"}},
    ],

    "discriminating": [
        ("HIGH", "H1 FY26 results (25 Feb): Australian Food EBIT growth rate is the single test. Consensus ~11.8% vs management \"mid-to-high single digits.\" A miss confirms H2 structural erosion thesis.",
         "H1 (turnaround on track) vs H2 (margin erosion structural)", "Awaiting: 15 days"),
        ("HIGH", "ACCC court proceedings outcome. Coles settled for $10.25M for similar conduct. Woolworths is fighting the case. Verdict sets precedent for penalties and remedies under mandatory FGCC.",
         "H1 (manageable) vs H3 (punitive precedent)", "Pending: date TBC"),
        ("HIGH", '$400M cost-out delivery rate. "On track" per management but unverified. H1 results will show first real evidence. If cost savings are consumed by price investment, H2 strengthens.',
         "H1 (margin recovery) vs H2 (savings consumed)", "On track per mgmt: unverified"),
        ("MED", "Aldi store expansion cadence. 602 stores. No online. Targeting premium segments. If Aldi launches online grocery, H4 accelerates materially.",
         "H1 (moat holds) vs H4 (moat breaks)", "No announced plans"),
    ],

    "non_discriminating": [
        "Revenue +1.7% FY25",
        "E-com +17% growth",
        "NZ recovery +38%",
        "Consensus Hold/Buy",
        "Target A$30.45",
    ],

    "tripwires": [
        {
            "date": "25 FEB 2026", "name": "H1 FY26 Results",
            "source": "ASX announcement, company earnings presentation, investor call",
            "conditions": [
                ("positive", "AU Food EBIT growth \u22658%", "H1 (Turnaround) strengthens materially. Demonstrates cost-out program is delivering margin recovery, not just funding price investment. Dividend recovery path confirmed."),
                ("negative", "AU Food EBIT growth <5%", 'H2 (Structural Erosion) strengthens. FY25 margin contraction was not a trough but a new normal. "Transitional" becomes "structural." Share price vulnerable; already above broker targets.'),
            ],
        },
        {
            "date": "H1 2026", "name": "ACCC v Woolworths",
            "source": "Federal Court; ACCC media releases",
            "conditions": [
                ("positive", "Settlement <$50M, no structural orders", "H3 (Regulatory) weakens. Manageable cost. Precedent set that mandatory FGCC enforcement is proportionate, not punitive."),
                ("negative", "Penalties >$100M or structural orders", "H3 crystallises. Sets precedent for aggressive enforcement under new code. Signals sustained regulatory pressure on business model."),
            ],
        },
        {
            "date": "APR 2026", "name": "Mandatory FGCC Impact",
            "source": "FGCC legislation; ACCC enforcement portal; supplier disclosures",
            "conditions": [
                ("positive", "Compliance absorbed within existing cost base", "Regulatory impact manageable. UK GCA precedent holds: moderate, not transformative."),
                ("negative", "Supplier renegotiation wave compresses gross margins", "Fresh produce provisions trigger structural change in supplier dynamics. H3 strengthens."),
            ],
        },
        {
            "date": "ONGOING", "name": "Aldi Online Launch",
            "source": "Aldi AU corporate communications; industry press; international market signals",
            "conditions": [
                ("positive", "Aldi remains offline", "Woolworths' e-commerce moat intact. 16%+ penetration is a structural advantage. H4 remains slow-burn."),
                ("negative", "Aldi launches e-grocery", "H4 accelerates from slow-burn to immediate. Woolworths' fastest-growing competitive advantage neutralised."),
            ],
        },
    ],

    "coverage": [
        ("Corporate Comms", "Full", "Q1 FY26", 'High, but motivated; "below ambition" caveat notable'),
        ("Regulatory Filings", "Full", "Apr 2025", "High: statutory source, mandatory code in force"),
        ("Broker Research", "Full", "Recent", "Medium: price above consensus target is a yellow flag"),
        ("Competitor Data", "Good", "Recent", "High: independent, verifiable"),
        ("Economic Data", "Good", "Jan 2026", "High: ABS, RBA sources"),
        ("Alternative Data", "Partial", "Mixed", "Medium: e-com data strong, workforce data limited"),
        ("Academic Research", "Limited", "Structural", "Medium: base rates, not predictions; UK precedent relevant but imperfect"),
        ("Media & Social", "Full", "Current", "Low: noise, but directionally consistent"),
    ],

    "questions": [
        "Can $400M cost-out absorb both price investment AND margin recovery, or is it one or the other?",
        "Is Aldi's lack of online a permanent strategic choice or a timing gap?",
        "Will mandatory FGCC penalties deter aggressive supplier negotiations, structurally changing gross margins?",
    ],

    "identity_table": [
        ("Ticker", "WOW.AX", "Exchange", "ASX"),
        ("Market Cap", "A$38.3B", "Enterprise Value", "A$54.1B"),
        ("Share Price", "A$31.41", "52-Week Range", "A$25.51 to A$33.76"),
        ("Forward P/E", "23.5x", "Trailing P/E", "36.6x"),
        ("EV/EBITDA", "10.5x", "Debt/Equity", "3.50x"),
        ("Dividend (FY25)", "A$0.84 (2.9%)", "Dividend Change", "Down 21%"),
        ("Revenue (FY25)", "A$69.1B (+1.7%)", "NPAT (FY25)", "A$963M (down 19.1%)"),
        ("EBIT Margin (FY25)", "4.0% (down 70bp)", "ROE", "18.1%"),
        ("Net Debt/EBITDA", "2.8x", "Free Cash Flow", "A$2.02B"),
        ("CEO", "Amanda Bardwell (since 2024)", "Employees", "202,000"),
        ("Last Earnings", "Aug 2025 (FY25)", "Next Earnings", "25 Feb 2026 (H1 FY26)"),
    ],

    "business_overview": (
        "Australia's largest supermarket operator with 38% grocery market share. "
        "Core Australian Food division contributes A$51.4B revenue. Also operates New Zealand "
        "supermarkets, Big W discount department stores, Petstock (acquired FY25), and B2B food "
        "distribution (PFD). E-commerce penetration reached 16.2% in Q1 FY26 with On Demand "
        "fulfilment growing 39%."
    ),

    "evidence_domains_long": [
        {
            "name": "1. Corporate Communications", "epistemic": "Motivated",
            "finding": (
                'New CEO Amanda Bardwell frames FY25 as "disappointing" but FY26 as "transitional." '
                "$400M cost-out program described as \"on track.\" Strategy centres on \"sharper value, "
                "more consistent availability, simplifying the business.\" Q1 FY26 trading update: "
                "AU Food sales +2.1% overall, +3.8% ex-tobacco. Customer NPS up 3 points, value "
                "perception up 5 points. E-commerce penetration at 16.2%. Medium-term target: "
                "mid-to-high single-digit EBIT growth, but acknowledged \"current trading performance "
                'has been below our ambition."'
            ),
            "tension": (
                'The admission that performance is "below ambition" is unusually candid for a CEO in her '
                "second year. Either this is honest expectation-setting (bullish signal: under-promise, "
                "over-deliver) or genuine uncertainty about whether the turnaround is working. The $400M "
                "cost-out was announced to fund price investment and margin recovery; it cannot do both "
                "simultaneously at scale. Which one wins?"
            ),
            "tags": [("Supports H1", "supports"), ("Neutral H2-H4", "neutral")],
            "source": "FY25 Results, Aug 2025; Q1 FY26 Trading Update; Investor presentations",
        },
        {
            "name": "2. Regulatory Filings & Proceedings", "epistemic": "Under Oath / Statutory",
            "finding": (
                "Three concurrent regulatory vectors: "
                "(a) ACCC Supermarkets Inquiry final report (March 2025): Confirmed oligopoly structure. "
                "20 recommendations for reform. Found both Coles and Woolworths increased product margins over five years. "
                "(b) ACCC v Woolworths (Federal Court): Proceedings commenced for allegedly misleading "
                '"Prices Dropped" claims on 266+ products. Woolworths is fighting (Coles settled). '
                "Penalties: up to the greater of $10M, 3x benefit, or 10% of turnover. "
                "(c) Mandatory Food and Grocery Code of Conduct took effect 1 April 2025. "
                "Significant penalties. Fresh produce provisions commence April 2026."
            ),
            "tension": (
                "The regulatory environment has fundamentally shifted from voluntary to mandatory, from monitoring "
                "to enforcement. Three separate actions running concurrently. The market has partially priced the "
                "headline risk but may be underestimating the cumulative compliance cost and behavioural constraints "
                "on supplier negotiations. Woolworths' decision to fight the ACCC (unlike Coles) adds litigation risk."
            ),
            "tags": [("Strongly Supports H3", "strong"), ("Supports H2", "supports"), ("Contradicts H1", "contradicts")],
            "source": "ACCC Final Report, Mar 2025; Federal Court proceedings; FGCC legislation",
        },
        {
            "name": "3. Broker Research", "epistemic": "Consensus",
            "finding": (
                "Consensus: Hold/Moderate Buy. 3 Buy, 8 Hold, 0 Sell across 11-15 covering analysts. "
                "Average 12-month target: A$30.02-30.45 (range: A$28.25 to A$33.70). Current price A$31.41 "
                "is above consensus target, implying the market is already pricing in more optimism than the "
                "sell-side. Consensus FY26 EPS: A$0.65 for H1. Consensus expects Australian Food EBIT growth "
                'of ~11.8%, materially above management\'s "mid-to-high single digit" guidance.'
            ),
            "tension": (
                "The share price trading above consensus target is unusual and significant. It means the market "
                "has either front-run the broker upgrades or is pricing in a more bullish scenario than the "
                "sell-side supports. If H1 FY26 misses the ~11.8% consensus growth expectation, there is no "
                "cushion in the broker targets; the price is already above them."
            ),
            "tags": [("Supports H1 (rating)", "supports"), ("Weakly Supports H2 (target below price)", "supports")],
            "source": "TipRanks, Investing.com, TradingView, MarketScreener consensus data",
        },
        {
            "name": "4. Competitor Disclosures", "epistemic": "Independent",
            "finding": (
                "Aldi Australia: 602 stores, A$13.3B+ annual sales, ~10% market share and growing. No online grocery "
                "offering. Targeting premium segments. Coles: 29% market share. Settled ACCC pricing case for $10.25M. "
                "Amazon: Pushing selected grocery categories, not yet full-service. Costco: Additional sites under consideration."
            ),
            "tension": (
                "Aldi's lack of online grocery is Woolworths' strongest competitive moat. If Aldi ever launches "
                "e-commerce (as it has in other markets), it would neutralise Woolworths' fastest-growing channel. "
                "Meanwhile, Woolworths' price cuts are a response to Aldi-driven competitive pressure; the cost-out "
                "program is partly funding defensive pricing, not margin expansion."
            ),
            "tags": [("Strongly Supports H4", "strong"), ("Supports H2", "supports")],
            "source": "ACCC Inquiry data; Aldi AU filings; Coles ASX announcements; industry reports",
        },
        {
            "name": "5. Economic Data", "epistemic": "Objective",
            "finding": (
                "Australian consumer sentiment recovering but still below historical averages. RBA rate cuts beginning "
                "to flow through. Grocery inflation moderating: prices ex-tobacco falling for seven consecutive quarters. "
                "Wage growth running ahead of CPI, creating cost pressure for labour-intensive retailers (202,000 employees). "
                "Tobacco sales declining materially; management estimates A$80-100M EBIT headwind in FY26."
            ),
            "tension": None,
            "tags": [("Weakly Supports H1", "supports"), ("Supports H2", "supports"), ("Supports H3", "supports")],
            "source": "ABS data; RBA; WOW Q1 FY26 Trading Update",
        },
        {
            "name": "6. Alternative Data", "epistemic": "Behavioural",
            "finding": (
                "E-commerce: On Demand orders +39%, 31% fulfilled within 2 hours. Auburn e-commerce CFC opened May 2025, "
                "Moorebank RDC opening H1 FY26. Customer: NPS +3 points, value perception +5 points YoY. "
                "Industrial relations: 17-day warehouse strike resolved with ~11% pay increase over 3 years. AI-driven "
                "performance monitoring was a key grievance; automation/workforce tension remains live."
            ),
            "tension": None,
            "tags": [("Supports H1 (e-com, NPS)", "supports"), ("Supports H4 (automation friction)", "supports")],
            "source": "WOW Q1 FY26 update; UWU disclosures; industry fulfilment data",
        },
        {
            "name": "7. Academic Research", "epistemic": "Peer-Reviewed",
            "finding": (
                "Research on grocery duopoly/oligopoly structures indicates price competition against discount entrants "
                "typically compresses incumbent margins by 100-200bp over a decade. E-commerce in grocery has negative "
                "margin characteristics in early scaling phases; fulfilment costs per order exceed in-store equivalents "
                "until very high penetration levels (typically >25%). UK Groceries Code Adjudicator precedent shows "
                "moderate ongoing compliance costs but limited impact on aggregate profitability."
            ),
            "tension": None,
            "tags": [("Supports H2 (margin compression)", "supports"), ("Supports H4 (discount dynamics)", "supports"),
                     ("Neutral on H3 (UK precedent moderate)", "neutral")],
            "source": "Published research on grocery duopoly dynamics; UK GCA impact studies",
        },
        {
            "name": "8. Media & Social", "epistemic": "Noise",
            "finding": (
                '"Price gouging" narrative remains politically salient. Woolworths previously among Australia\'s most '
                "trusted brands; trust has deteriorated sharply. Cost-of-living angle makes supermarkets a political "
                'target. CEO Bardwell conceded group performance is "not good enough." Media framing consistently '
                "negative. Share price down 11.6% over five years."
            ),
            "tension": None,
            "tags": [("Contradicts H1", "contradicts"), ("Supports H2", "supports"), ("Strongly Supports H3", "strong")],
            "source": "Major Australian financial press; social sentiment tracking; brand trust surveys",
        },
    ],

    "non_discrim_long": (
        'The following are true but do not distinguish between hypotheses: "Revenue +1.7% FY25" '
        "(low growth consistent with all scenarios). \"E-commerce +17%\" (impressive but margin impact "
        "unclear; could support H1 or H2 depending on unit economics). \"NZ recovery +38%\" (small division, "
        "doesn't move group thesis). \"Consensus Hold/Buy\" (price already above target; consensus is not "
        'providing upside signal). "Analyst target A$30.45" (below current price; arguably supports H2).'
    ),

    "gaps": [
        ("E-commerce unit economics", "Woolworths does not disclose e-commerce profitability separately. Cannot assess whether 16%+ penetration is margin-accretive or dilutive. This is the single biggest gap for evaluating H1 vs H2."),
        ("$400M cost-out granularity", 'Management says "on track" but has not disclosed category-level savings or how much is being reinvested into pricing vs flowing to EBIT.'),
        ("Mandatory FGCC compliance cost", "The code only took effect April 2025. No company has yet disclosed the actual compliance cost burden. True financial impact is not yet knowable."),
        ("Industrial relations trajectory", "The 17-day strike was resolved with ~11% pay increases. Whether this settlement sets a precedent for future enterprise bargaining is unknown."),
    ],

    "limitations": (
        "Hypothesis survival scores (H1: 50%, H2: 35%, H3: 30%, H4: 25%) reflect editorial assessment based "
        "on publicly available evidence. These are qualitative judgments, not algorithmic outputs. The scores "
        "sum to more than 100% because H3 (Regulatory) is partially independent of H1/H2; regulatory outcomes "
        "can compound with either turnaround or erosion scenarios. The formal scoring methodology is under development."
    ),

    "disclaimer": (
        "This report does not constitute personal financial advice under Australian law or any other jurisdiction. "
        "Continuum Trinity synthesises cross-domain evidence using the Analysis of Competing Hypotheses (ACH) methodology. "
        "All factual claims are sourced from ASX filings, ACCC publications, broker consensus data, and publicly available "
        "information as at the report date. Hypothesis survival scores reflect editorial assessment. This report does not "
        "contain buy, sell, or hold recommendations, price targets, or valuation models. Readers should consult their own "
        "financial advisors before making investment decisions."
    ),
}


# ═══════════════════════════════════════════════════════════════
# SIGNAL RENDERING
# ═══════════════════════════════════════════════════════════════

SIGNAL_MAP = {
    "strong_support": ("\u25B2\u25B2", C["green"]),
    "support":        ("\u25B2",       C["green"]),
    "weak_support":   ("\u25B2",       RGBColor(0x68, 0xD3, 0x91)),
    "contradict":     ("\u25BC",       C["red"]),
    "neutral":        ("\u2014",       C["text_muted"]),
}

TAG_COLOURS = {
    "supports":   ("E8F5E9", RGBColor(0x2E, 0x7D, 0x32)),
    "contradicts": ("FFEBEE", RGBColor(0xC6, 0x28, 0x28)),
    "neutral":    ("F5F5F5", RGBColor(0x9E, 0x9E, 0x9E)),
    "strong":     ("C8E6C9", RGBColor(0x1B, 0x5E, 0x20)),
}


# ═══════════════════════════════════════════════════════════════
# SHORT FORM BUILDER
# ═══════════════════════════════════════════════════════════════

def build_short_form(doc):
    d = DATA

    # Page setup: A4 portrait, tight margins
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9)
    style.font.color.rgb = C["charcoal"]
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # ── HEADER BAR ──
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl.rows[0].cells:
        set_cell_shading(cell, HEX["midnight"])
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    # Left: ticker + company + sector
    left = tbl.cell(0, 0)
    p = left.paragraphs[0]
    add_run(p, d["ticker"], size=16, colour=C["white"], bold=True)
    add_run(p, f"   {d['company']}", size=9, colour=C["text_muted"])

    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    add_badge_run(p2, f"{d['sector']}  \u2022  {d['subsector']}", "1A5F6C", C["sage"], size=6)

    # Right: brand + date
    right = tbl.cell(0, 1)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = add_run(p, "CONTINUUM ", size=7, colour=C["sage"], bold=True)
    add_run(p, "TRINITY", size=7, colour=C["text_muted"], bold=True)
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    add_run(p2, f"{d['date']}  \u2022  {d['version']}", size=6, colour=C["text_muted"], font=MONO)

    # ── PRICE METRICS BAR ──
    spacer(doc, 2)
    ncols = len(d["metrics"]) + 1
    tbl = doc.add_table(rows=2, cols=ncols)
    remove_table_borders(tbl)
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_shading(cell, HEX["sidebar"])
            set_cell_margins(cell, top=40, bottom=40, left=60, right=60)

    # Price cell spans 2 rows
    price_cell = tbl.cell(0, 0)
    price_cell.merge(tbl.cell(1, 0))
    set_cell_vertical_alignment(price_cell, "center")
    p = price_cell.paragraphs[0]
    add_run(p, "A$", size=10, colour=C["text_muted"], font=MONO)
    add_run(p, "31.41", size=18, colour=C["text_primary"], bold=True, font=MONO)

    colour_map = {"premium": C["amber"], "negative": C["red"], "positive": C["green"]}

    for i, (label, value, ctype) in enumerate(d["metrics"]):
        # Label row
        cell_label = tbl.cell(0, i + 1)
        p = cell_label.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, label.upper(), size=5.5, colour=C["text_muted"], bold=True)

        # Value row
        cell_val = tbl.cell(1, i + 1)
        p = cell_val.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        val_colour = colour_map.get(ctype, C["text_secondary"])
        add_run(p, value, size=8, colour=val_colour, font=MONO, bold=(ctype != ""))

    # ── RISK SKEW BAR ──
    spacer(doc, 2)
    tbl = doc.add_table(rows=1, cols=3)
    remove_table_borders(tbl)
    for cell in tbl.rows[0].cells:
        set_cell_margins(cell, top=40, bottom=40, left=60, right=60)

    # Label
    label_cell = tbl.cell(0, 0)
    p = label_cell.paragraphs[0]
    add_run(p, "RISK SKEW", size=6, colour=C["text_muted"], bold=True)
    tc = label_cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="700" w:type="dxa"/>')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

    # Badge
    skew = d["risk_skew"]
    skew_colours = {
        "Downside": ("FFEBEE", C["red"], "\u25BC"),
        "Upside":   ("E8F5E9", C["green"], "\u25B2"),
        "Balanced": ("FFF3E0", C["amber"], "\u25C6"),
    }
    bg_hex, skew_col, arrow = skew_colours.get(skew, skew_colours["Balanced"])

    badge_cell = tbl.cell(0, 1)
    p = badge_cell.paragraphs[0]
    add_badge_run(p, f"{arrow} {skew.upper()}", bg_hex, skew_col, size=8)
    tc = badge_cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="1200" w:type="dxa"/>')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

    # Rationale
    rat_cell = tbl.cell(0, 2)
    p = rat_cell.paragraphs[0]
    add_run(p, d["risk_skew_rationale"], size=7, colour=C["text_secondary"])

    # ── DOMINANT NARRATIVE ──
    section_label(doc, "Dominant Narrative")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, d["narrative"], size=8.5, colour=C["text_secondary"])

    # Verdict callout
    p = doc.add_paragraph()
    set_para_border_left(p, "C07A1A", sz=18, space=8)
    set_para_shading(p, HEX["callout_warn"])
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Mm(2)
    add_run(p, d["verdict"], size=8, colour=C["amber"], bold=True)

    # ── HYPOTHESIS SURVIVAL BAR ──
    section_label(doc, "Hypothesis Survival")
    hyps = d["hypotheses"]
    total = sum(h["score"] for h in hyps)

    tbl = doc.add_table(rows=1, cols=len(hyps))
    remove_table_borders(tbl)
    set_row_height(tbl.rows[0], 14)

    for i, h in enumerate(hyps):
        cell = tbl.cell(0, i)
        _, hex_c = HYP_COLOURS[h["id"]]
        set_cell_shading(cell, hex_c)
        set_cell_margins(cell, top=20, bottom=20, left=20, right=20)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"{h['id']}: {h['score']}%", size=6.5, colour=C["white"], bold=True, font=MONO)
        # Set proportional width
        pct = int(5000 * h["score"] / total)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{pct}" w:type="pct"/>')
        existing = tcPr.find(qn('w:tcW'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcW)

    # Legend
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    for h in hyps:
        rgb, hex_c = HYP_COLOURS[h["id"]]
        add_run(p, "\u25CF ", size=7, colour=rgb)
        add_run(p, f"{h['name']}   ", size=6, colour=C["text_muted"])

    # ── HYPOTHESIS CARDS ──
    for h in hyps:
        rgb, hex_c = HYP_COLOURS[h["id"]]
        tbl = doc.add_table(rows=1, cols=2)
        remove_table_borders(tbl)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Left colour strip
        strip = tbl.cell(0, 0)
        set_cell_shading(strip, hex_c)
        set_cell_width(strip, Mm(3))
        set_cell_margins(strip, left=0, right=0)
        strip.paragraphs[0].add_run("").font.size = Pt(1)

        # Set strip column width
        tc = strip._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="150" w:type="dxa"/>')
        existing = tcPr.find(qn('w:tcW'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcW)

        # Content cell
        content = tbl.cell(0, 1)
        set_cell_shading(content, HEX["sidebar"])
        set_cell_margins(content, top=50, bottom=50, left=100, right=100)

        p = content.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        add_run(p, f"{h['id']}: {h['name']}", size=8, colour=C["text_primary"], bold=True)
        add_run(p, "  ", size=7)

        # Status badge
        status_badges = {
            "Priced": ("E8F5E9", C["green"]),
            "Evidence Building": ("FFF3E0", C["amber"]),
            "Active": ("FFEBEE", C["red"]),
            "Slow Burn": ("F5F5F5", C["text_muted"]),
        }
        bg, tc_col = status_badges.get(h["status"], ("F5F5F5", C["text_muted"]))
        add_badge_run(p, h["status"].upper(), bg, tc_col, size=5.5)

        # Score + direction right-aligned
        p2 = content.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        add_run(p2, h["desc"], size=7, colour=C["text_muted"])

        p3 = content.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.paragraph_format.space_before = Pt(1)
        p3.paragraph_format.space_after = Pt(0)
        add_run(p3, f"{h['score']}%", size=12, colour=rgb, bold=True, font=MONO)
        dir_colours = {"Rising": C["red"], "Awaiting": C["text_muted"], "Steady": C["text_muted"]}
        dir_arrows = {"Rising": "\u2191 ", "Awaiting": "\u2192 ", "Steady": "\u2192 "}
        add_run(p3, f"  {dir_arrows.get(h['direction'], '')}{h['direction']}", size=7,
                colour=dir_colours.get(h["direction"], C["text_muted"]), bold=True)

        spacer(doc, 2)

    # ── EVIDENCE MATRIX ──
    section_label(doc, "Cross-Domain Evidence Matrix")
    hyp_ids = ["H1", "H2", "H3", "H4"]
    hyp_names = {h["id"]: h["name"].split()[0] for h in hyps}  # Short names
    full_names = ["H1 Turnaround", "H2 Erosion", "H3 Regulatory", "H4 Disruption"]

    matrix = d["evidence_matrix"]
    tbl = doc.add_table(rows=len(matrix) + 1, cols=5)
    set_table_borders(tbl, "E2E8F0", 4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Domain"] + full_names
    for i, hdr in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, HEX["table_hdr"])
        set_cell_margins(cell, top=30, bottom=30, left=50, right=50)
        p = cell.paragraphs[0]
        if i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, hdr.upper(), size=6, colour=C["white"], bold=True)

    # Data rows
    for r, ev in enumerate(matrix):
        for c in range(5):
            cell = tbl.cell(r + 1, c)
            set_cell_margins(cell, top=20, bottom=20, left=50, right=50)
            if r % 2 == 1:
                set_cell_shading(cell, HEX["alt_row"])
            p = cell.paragraphs[0]
            if c == 0:
                add_run(p, ev["domain"], size=7.5, colour=C["text_primary"], bold=True)
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
                add_run(p2, ev["epistemic"], size=6, colour=C["text_muted"], italic=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hid = hyp_ids[c - 1]
                sig = ev["signals"][hid]
                symbol, colour = SIGNAL_MAP.get(sig, ("\u2014", C["text_muted"]))
                add_run(p, symbol, size=10, colour=colour, bold=True)

    # ── WHAT DISCRIMINATES ──
    section_label(doc, "What Discriminates")
    for diag, text, _, _ in d["discriminating"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Mm(1)
        diag_bg = "FFEBEE" if diag == "HIGH" else "FFF3E0"
        diag_col = C["red"] if diag == "HIGH" else C["amber"]
        add_badge_run(p, diag, diag_bg, diag_col, size=6)
        add_run(p, f"  {text}", size=7.5, colour=C["text_secondary"])

    # Non-discriminating
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "ASSESSED & DISCARDED (NON-DISCRIMINATING):  ", size=6, colour=C["text_muted"], bold=True)
    for i, nd in enumerate(d["non_discriminating"]):
        add_run(p, nd, size=7, colour=C["text_muted"], strike=True)
        if i < len(d["non_discriminating"]) - 1:
            add_run(p, "   |   ", size=7, colour=C["text_muted"])

    # ── TRIPWIRES ──
    section_label(doc, "What We're Watching")
    for tw in d["tripwires"]:
        tbl = doc.add_table(rows=1 + len(tw["conditions"]), cols=2)
        set_table_borders(tbl, "E2E8F0", 4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        date_cell = tbl.cell(0, 0)
        name_cell = tbl.cell(0, 1)
        for c in [date_cell, name_cell]:
            set_cell_margins(c, top=30, bottom=30, left=80, right=80)
        p = date_cell.paragraphs[0]
        add_run(p, tw["date"], size=8, colour=C["gold"], bold=True, font=MONO)
        p = name_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, tw["name"].upper(), size=6.5, colour=C["text_muted"], bold=True)

        # Condition rows
        for ci, (direction, condition, consequence) in enumerate(tw["conditions"]):
            # Merge into single cell for each condition
            cell = tbl.cell(ci + 1, 0)
            cell.merge(tbl.cell(ci + 1, 1))
            set_cell_margins(cell, top=20, bottom=20, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            arrow = "\u25B2 " if direction == "positive" else "\u25BC "
            arrow_col = C["green"] if direction == "positive" else C["red"]
            add_run(p, arrow, size=8, colour=arrow_col, bold=True)
            add_run(p, condition, size=7.5, colour=C["text_secondary"])

        spacer(doc, 3)

    # ── EVIDENCE COVERAGE ──
    section_label(doc, "Evidence Coverage")
    tbl = doc.add_table(rows=len(d["coverage"]) + 1, cols=2)
    set_table_borders(tbl, "E2E8F0", 4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for c in range(2):
        cell = tbl.cell(0, c)
        set_cell_shading(cell, HEX["table_hdr"])
        set_cell_margins(cell, top=25, bottom=25, left=80, right=80)
    add_run(tbl.cell(0, 0).paragraphs[0], "DOMAIN", size=6.5, colour=C["white"], bold=True)
    p = tbl.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "COVERAGE", size=6.5, colour=C["white"], bold=True)

    for i, (domain, coverage, _, _) in enumerate(d["coverage"]):
        cell_d = tbl.cell(i + 1, 0)
        set_cell_margins(cell_d, top=15, bottom=15, left=80, right=80)
        add_run(cell_d.paragraphs[0], domain, size=8, colour=C["text_secondary"])

        cell_c = tbl.cell(i + 1, 1)
        set_cell_margins(cell_c, top=15, bottom=15, left=80, right=80)
        p = cell_c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg, tc_col = COV_BADGES.get(coverage, ("F5F5F5", C["text_muted"]))
        add_badge_run(p, coverage.upper(), bg, tc_col, size=6.5)

    # ── UNANSWERED QUESTIONS ──
    section_label(doc, "Unanswered Questions")
    for q in d["questions"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Mm(3)
        p.paragraph_format.first_line_indent = Mm(-3)
        add_run(p, "?  ", size=8, colour=C["deep_teal"], bold=True, font=MONO)
        add_run(p, q, size=7.5, colour=C["text_secondary"])

    # ── FOOTER ──
    spacer(doc, 4)
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    for cell in tbl.rows[0].cells:
        set_cell_shading(cell, HEX["midnight"])
        set_cell_margins(cell, top=60, bottom=60, left=120, right=120)

    p = tbl.cell(0, 0).paragraphs[0]
    add_run(p, d["disclaimer"][:200] + "...", size=5.5, colour=C["text_muted"], italic=True)

    p = tbl.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"ID: {d['report_id']}", size=6, colour=C["text_muted"], font=MONO)
    p2 = tbl.cell(0, 1).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    add_run(p2, f"MODE: {d['mode']}", size=6, colour=C["text_muted"], font=MONO)
    p3 = tbl.cell(0, 1).add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(0)
    add_run(p3, f"NEXT: {d['next_update']}", size=6, colour=C["text_muted"], font=MONO)


# ═══════════════════════════════════════════════════════════════
# LONG FORM BUILDER
# ═══════════════════════════════════════════════════════════════

def build_long_form(doc):
    d = DATA

    # Page setup
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(9)
    style.font.color.rgb = C["charcoal"]
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(14)

    # Running header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(hp, "CONTINUUM ", size=7, colour=C["sage"], bold=True)
    add_run(hp, "TRINITY", size=7, colour=C["text_muted"], bold=True)
    add_run(hp, f"   |   {d['ticker']}", size=7, colour=C["text_muted"], font=MONO)

    # Running footer
    footer = section.footer
    fp = footer.paragraphs[0]
    add_run(fp, f"Confidential   |   {d['report_id']}   |   {d['date']}", size=6, colour=C["text_muted"])

    # ══ COVER HEADER ══
    tbl = doc.add_table(rows=3, cols=2)
    remove_table_borders(tbl)
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_shading(cell, HEX["midnight"])
            set_cell_margins(cell, top=40, bottom=40, left=150, right=150)

    # Row 0: brand + date
    p = tbl.cell(0, 0).paragraphs[0]
    add_run(p, "CONTINUUM ", size=7, colour=C["sage"], bold=True)
    add_run(p, "TRINITY", size=7, colour=C["text_muted"], bold=True)
    p = tbl.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, f"{d['date']}  \u2022  Investment Report {d['version']}", size=7, colour=C["text_muted"], font=MONO)

    # Row 1: title + price
    p = tbl.cell(1, 0).paragraphs[0]
    add_run(p, "NARRATIVE INTELLIGENCE", size=6, colour=C["sage"], bold=True)
    p2 = tbl.cell(1, 0).add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    add_run(p2, "Woolworths Group", size=22, colour=C["white"], bold=True)
    p3 = tbl.cell(1, 0).add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    add_run(p3, f"{d['ticker']}  \u2022  ASX  \u2022  {d['sector']}", size=9, colour=C["text_muted"])
    p4 = tbl.cell(1, 0).add_paragraph()
    p4.paragraph_format.space_before = Pt(4)
    add_badge_run(p4, "38% Grocery Market Share  \u2022  202,000 Employees  \u2022  A$69.1B Revenue", "1A5F6C", C["sage"], size=6)

    # Price + metrics in right cell
    p = tbl.cell(1, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, "A$", size=10, colour=C["text_muted"], font=MONO)
    add_run(p, "31.41", size=22, colour=C["white"], bold=True, font=MONO)

    # Key metrics under price
    key_metrics = [("Mkt Cap", "A$38.3B", ""), ("Fwd P/E", "23.5x", "premium"),
                   ("EV/EBITDA", "10.5x", ""), ("NPAT FY25", "\u219319%", "negative"),
                   ("Div Yield", "2.9%", "")]
    colour_map = {"premium": C["gold"], "negative": C["red_light"], "positive": C["light_green"]}
    p2 = tbl.cell(1, 1).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(6)
    for label, val, ctype in key_metrics:
        add_run(p2, f"{label.upper()}: ", size=5.5, colour=C["text_muted"])
        val_col = colour_map.get(ctype, C["white"])
        add_run(p2, f"{val}  ", size=8, colour=val_col, font=MONO, bold=bool(ctype))

    # Row 2: subtitle bar
    tbl.cell(2, 0).merge(tbl.cell(2, 1))
    merged = tbl.cell(2, 0)
    set_cell_shading(merged, HEX["midnight"])
    # Add bottom border effect via sage line
    p = merged.paragraphs[0]
    set_para_border_bottom(p, "4A9E7E", sz=12, space=0)

    # ══ VERDICT BAR ══
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    for cell in tbl.rows[0].cells:
        set_cell_shading(cell, HEX["midnight"])
        set_cell_margins(cell, top=60, bottom=60, left=150, right=150)

    p = tbl.cell(0, 0).paragraphs[0]
    add_run(p, d["verdict_long"], size=9.5, colour=C["gold"], bold=True)

    # Risk Skew line in verdict bar
    skew = d["risk_skew"]
    skew_colours_long = {
        "Downside": (C["red_light"], "\u25BC"),
        "Upside":   (C["light_green"], "\u25B2"),
        "Balanced": (C["gold"], "\u25C6"),
    }
    skew_col, skew_arrow = skew_colours_long.get(skew, (C["gold"], "\u25C6"))
    p_skew = tbl.cell(0, 0).add_paragraph()
    p_skew.paragraph_format.space_before = Pt(6)
    p_skew.paragraph_format.space_after = Pt(0)
    add_run(p_skew, "RISK SKEW   ", size=6, colour=C["text_muted"])
    add_run(p_skew, f"{skew_arrow} {skew.upper()}", size=9, colour=skew_col, bold=True, font=MONO)

    p = tbl.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for h in d["hypotheses"]:
        rgb, _ = HYP_COLOURS[h["id"]]
        add_run(p, f"{h['id']} ", size=6, colour=C["text_muted"])
        add_run(p, f"{h['score']}% ", size=11, colour=rgb, bold=True, font=MONO)
        dir_col = C["red_light"] if h["direction"] == "Rising" else C["text_muted"]
        dir_arrow = "\u2191" if h["direction"] == "Rising" else "\u2192"
        add_run(p, f"{dir_arrow} ", size=7, colour=dir_col)
        add_run(p, "  ", size=6)

    spacer(doc, 6)

    # ══ SECTION 1: IDENTITY & SNAPSHOT ══
    _section_heading(doc, 1, "Identity & Snapshot")

    id_tbl = doc.add_table(rows=len(d["identity_table"]), cols=4)
    set_table_borders(id_tbl, "E2E8F0", 4)
    for r, (l1, v1, l2, v2) in enumerate(d["identity_table"]):
        for ci, (label, val) in enumerate([(l1, v1), (l2, v2)]):
            lc = id_tbl.cell(r, ci * 2)
            vc = id_tbl.cell(r, ci * 2 + 1)
            set_cell_margins(lc, top=15, bottom=15, left=60, right=20)
            set_cell_margins(vc, top=15, bottom=15, left=20, right=60)
            if r % 2 == 0:
                set_cell_shading(lc, HEX["alt_row"])
                set_cell_shading(vc, HEX["alt_row"])
            add_run(lc.paragraphs[0], label, size=8, colour=C["text_primary"], bold=True)
            add_run(vc.paragraphs[0], val, size=8, colour=C["text_secondary"], font=MONO)

    spacer(doc, 2)
    p = doc.add_paragraph()
    add_run(p, "Business overview: ", size=8.5, colour=C["text_primary"], bold=True)
    add_run(p, d["business_overview"], size=8.5, colour=C["text_secondary"])

    # ══ SECTION 2: DOMINANT NARRATIVE ══
    doc.add_page_break()
    _section_heading(doc, 2, "Dominant Narrative")

    _subsection(doc, "The Narrative")
    p = doc.add_paragraph()
    add_run(p, d["narrative"], size=8.5, colour=C["text_secondary"])

    _subsection(doc, "The Price Implication")
    _callout(doc, "Embedded Assumptions at A$31.41",
             "Australian Food EBIT recovery: Mid-to-high single-digit growth in FY26. "
             "Cost-out delivery: $400M savings realised by end CY25. "
             "E-commerce scales profitably: 16%+ penetration without margin dilution. "
             "Regulatory manageable: Mandatory FGCC compliance costs absorbed. "
             "Competitive position holds: Aldi stays offline. "
             "Big W returns to profit: Currently loss-making at -$63M EBIT.",
             "teal")

    _subsection(doc, "The Evidence Check")
    p = doc.add_paragraph()
    add_run(p, (
        "The turnaround narrative has supporting evidence (Q1 FY26 trading data, cost-out program on track, "
        "NPS improvement) but also faces material headwinds that are structural, not cyclical: mandatory "
        "regulatory code now in force, ACCC court proceedings pending, Aldi expanding, consumer price sensitivity "
        'elevated, and management itself has called FY26 "transitional" with recovery "below ambition."'
    ), size=8.5, colour=C["text_secondary"])

    _subsection(doc, "Narrative Stability")
    p = doc.add_paragraph()
    add_run(p, "Fragile. ", size=8.5, colour=C["text_primary"], bold=True)
    add_run(p, (
        "The turnaround narrative is new; Bardwell has been CEO for approximately 18 months. "
        "FY25 delivered a 19% NPAT decline. The narrative rests almost entirely on forward promises "
        "($400M cost-out, FY26 as transitional) rather than demonstrated results. The H1 FY26 results "
        "on 25 February will be the first real test."
    ), size=8.5, colour=C["text_secondary"])

    # ══ SECTION 3: CROSS-DOMAIN EVIDENCE SYNTHESIS ══
    doc.add_page_break()
    _section_heading(doc, 3, "Cross-Domain Evidence Synthesis")
    p = doc.add_paragraph()
    add_run(p, "Eight evidence domains assessed. Evidence ranked: ", size=8.5, colour=C["text_secondary"])
    add_run(p, "Facts", size=8.5, colour=C["text_primary"], bold=True)
    add_run(p, " (filed, audited) > ", size=8.5, colour=C["text_secondary"])
    add_run(p, "Company Releases", size=8.5, colour=C["text_primary"], bold=True)
    add_run(p, " (motivated) > ", size=8.5, colour=C["text_secondary"])
    add_run(p, "Broker Research", size=8.5, colour=C["text_primary"], bold=True)
    add_run(p, " (consensus) > ", size=8.5, colour=C["text_secondary"])
    add_run(p, "Social & Media", size=8.5, colour=C["text_primary"], bold=True)
    add_run(p, " (noise).", size=8.5, colour=C["text_secondary"])

    for ed in d["evidence_domains_long"]:
        spacer(doc, 4)
        _evidence_domain_card(doc, ed)

    # Evidence Alignment Summary table
    spacer(doc, 4)
    _subsection(doc, "Evidence Alignment Summary")
    hyp_ids = ["H1", "H2", "H3", "H4"]
    full_names = ["H1 Turnaround", "H2 Erosion", "H3 Regulatory", "H4 Disruption"]
    matrix = d["evidence_matrix"]

    tbl = doc.add_table(rows=len(matrix) + 2, cols=6)
    set_table_borders(tbl, "E2E8F0", 4)
    headers = ["Domain", "Epistemic Status"] + full_names
    for i, hdr in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, HEX["table_hdr"])
        set_cell_margins(cell, top=25, bottom=25, left=40, right=40)
        p = cell.paragraphs[0]
        if i >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, hdr.upper(), size=6, colour=C["white"], bold=True)

    for r, ev in enumerate(matrix):
        if r % 2 == 1:
            for c in range(6):
                set_cell_shading(tbl.cell(r + 1, c), HEX["alt_row"])
        for c in range(6):
            cell = tbl.cell(r + 1, c)
            set_cell_margins(cell, top=15, bottom=15, left=40, right=40)
            p = cell.paragraphs[0]
            if c == 0:
                add_run(p, ev["domain"], size=7.5, colour=C["text_primary"], bold=True)
            elif c == 1:
                add_run(p, ev["epistemic"], size=7, colour=C["text_muted"], italic=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sig = ev["signals"][hyp_ids[c - 2]]
                symbol, colour = SIGNAL_MAP.get(sig, ("\u2014", C["text_muted"]))
                add_run(p, symbol, size=9, colour=colour, bold=True)

    # Summary row
    summary_row = len(matrix) + 1
    for c in range(6):
        cell = tbl.cell(summary_row, c)
        set_cell_margins(cell, top=15, bottom=15, left=40, right=40)
    add_run(tbl.cell(summary_row, 0).paragraphs[0], "Domain Count", size=7.5, colour=C["text_primary"], bold=True)
    add_run(tbl.cell(summary_row, 1).paragraphs[0], "", size=7)
    counts = ["3-4 (mixed)", "5", "3 (2 strong)", "3"]
    count_colours = [C["text_muted"], C["amber"], C["red"], C["text_muted"]]
    for i, (ct, cc) in enumerate(zip(counts, count_colours)):
        p = tbl.cell(summary_row, i + 2).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, ct, size=7.5, colour=cc, bold=True, font=MONO)

    # ══ SECTION 4: COMPETING HYPOTHESES ══
    doc.add_page_break()
    _section_heading(doc, 4, "Competing Hypotheses")

    for h in d["hypotheses"]:
        _hypothesis_card_long(doc, h)
        spacer(doc, 4)

    # ══ SECTION 5: WHAT DISCRIMINATES ══
    doc.add_page_break()
    _section_heading(doc, 5, "What Discriminates")
    p = doc.add_paragraph()
    add_run(p, "Four data points carry high diagnosticity. Everything else is either confirmatory or consistent with multiple hypotheses.", size=8.5, colour=C["text_secondary"])

    # Discriminating evidence table
    spacer(doc, 2)
    tbl = doc.add_table(rows=len(d["discriminating"]) + 1, cols=4)
    set_table_borders(tbl, "E2E8F0", 4)
    for i, hdr in enumerate(["Diagnosticity", "Evidence", "Discriminates Between", "Current Reading"]):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, HEX["table_hdr"])
        set_cell_margins(cell, top=25, bottom=25, left=50, right=50)
        add_run(cell.paragraphs[0], hdr.upper(), size=6, colour=C["white"], bold=True)

    for r, (diag, text, between, reading) in enumerate(d["discriminating"]):
        for c in range(4):
            cell = tbl.cell(r + 1, c)
            set_cell_margins(cell, top=20, bottom=20, left=50, right=50)
            p = cell.paragraphs[0]
            if c == 0:
                diag_col = C["red"] if diag == "HIGH" else C["amber"]
                add_run(p, diag, size=8, colour=diag_col, bold=True)
            elif c == 1:
                add_run(p, text, size=7.5, colour=C["text_secondary"])
            elif c == 2:
                add_run(p, between, size=7.5, colour=C["text_secondary"])
            else:
                reading_col = C["amber"] if "Awaiting" in reading or "Pending" in reading or "unverified" in reading else C["green"]
                add_run(p, reading, size=7.5, colour=reading_col)

    spacer(doc, 4)
    _callout(doc, "Non-Discriminating Evidence: Assessed & Discarded", d["non_discrim_long"], "warn")

    # ══ SECTION 6: WHAT'S CHANGED ══
    _section_heading(doc, 6, "What's Changed")
    _callout(doc, "Initial Report",
             "This is the first Investment Report for Woolworths Group. No prior reference point. "
             "The evidence state as of 10 February 2026 establishes the baseline. Key evidence to track: "
             "H1 FY26 results (25 Feb), ACCC court proceedings timeline, mandatory FGCC first enforcement "
             "actions, Aldi store openings and any online announcements.",
             "teal")

    # ══ SECTION 7: WHAT WE'RE WATCHING ══
    _section_heading(doc, 7, "What We're Watching")
    p = doc.add_paragraph()
    add_run(p, "Four revision conditions with specific thresholds. Defined before the events occur to reduce rationalisation bias.", size=8.5, colour=C["text_secondary"])

    for tw in d["tripwires"]:
        spacer(doc, 4)
        _tripwire_card_long(doc, tw)

    # ══ SECTION 8: EVIDENCE GAPS & INTEGRITY ══
    doc.add_page_break()
    _section_heading(doc, 8, "Evidence Gaps & Integrity Notes")

    _subsection(doc, "Domain Coverage Assessment")
    tbl = doc.add_table(rows=len(d["coverage"]) + 1, cols=4)
    set_table_borders(tbl, "E2E8F0", 4)
    for i, hdr in enumerate(["Domain", "Coverage", "Freshness", "Confidence"]):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, HEX["table_hdr"])
        set_cell_margins(cell, top=25, bottom=25, left=50, right=50)
        add_run(cell.paragraphs[0], hdr.upper(), size=6, colour=C["white"], bold=True)

    for r, (domain, coverage, freshness, confidence) in enumerate(d["coverage"]):
        for c in range(4):
            cell = tbl.cell(r + 1, c)
            set_cell_margins(cell, top=15, bottom=15, left=50, right=50)
            if r % 2 == 1:
                set_cell_shading(cell, HEX["alt_row"])
            p = cell.paragraphs[0]
            if c == 0:
                add_run(p, domain, size=8, colour=C["text_primary"], bold=True)
            elif c == 1:
                bg, tc_col = COV_BADGES.get(coverage, ("F5F5F5", C["text_muted"]))
                # Coverage dot
                dot_colours = {"Full": C["green"], "Good": C["sage"], "Partial": C["amber"], "Limited": C["text_muted"]}
                add_run(p, "\u25CF ", size=8, colour=dot_colours.get(coverage, C["text_muted"]))
                add_run(p, coverage, size=8, colour=C["text_secondary"])
            elif c == 2:
                add_run(p, freshness, size=7.5, colour=C["text_muted"], font=MONO)
            else:
                # Colour confidence by level
                conf_col = C["green"] if "High" in confidence else (C["amber"] if "Medium" in confidence else C["text_muted"])
                add_run(p, confidence, size=7, colour=conf_col)

    _subsection(doc, "What We Couldn't Assess")
    for title, desc in d["gaps"]:
        _callout(doc, title, desc, "teal")

    _subsection(doc, "Analytical Limitations")
    p = doc.add_paragraph()
    add_run(p, d["limitations"], size=8.5, colour=C["text_secondary"])

    # ══ FOOTER ══
    spacer(doc, 8)
    tbl = doc.add_table(rows=2, cols=1)
    remove_table_borders(tbl)
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_shading(cell, HEX["midnight"])
            set_cell_margins(cell, top=40, bottom=40, left=150, right=150)

    p = tbl.cell(0, 0).paragraphs[0]
    add_run(p, d["disclaimer"], size=6.5, colour=C["text_muted"], italic=True)

    p = tbl.cell(1, 0).paragraphs[0]
    r = add_run(p, "CONTINUUM ", size=6.5, colour=C["sage"], bold=True)
    add_run(p, "TRINITY", size=6.5, colour=C["text_muted"], bold=True)
    add_run(p, f"      ID: {d['report_id']}   |   Mode: {d['mode']}   |   "
            f"Domains: 8 of 8   |   Hypotheses: 4 Active   |   Next: {d['next_update']}",
            size=6, colour=C["text_muted"], font=MONO)


# ═══════════════════════════════════════════════════════════════
# LONG FORM COMPONENT HELPERS
# ═══════════════════════════════════════════════════════════════

def _section_heading(doc, num, title):
    spacer(doc, 6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"SECTION {num:02d}", size=7, colour=C["deep_teal"], bold=True, font=MONO)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_para_border_bottom(p, "1A5F6C", sz=8, space=4)
    add_run(p, title, size=14, colour=C["text_primary"], bold=True)


def _subsection(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, size=10, colour=C["deep_teal"], bold=True)


def _callout(doc, label, text, variant="teal"):
    colours = {
        "teal":     (HEX["callout"],      "1A5F6C", C["deep_teal"]),
        "warn":     (HEX["callout_warn"], "C07A1A", C["amber"]),
        "critical": (HEX["callout_red"],  "C53030", C["red"]),
        "positive": (HEX["callout_green"], "2F855A", C["green"]),
    }
    bg_hex, border_hex, label_colour = colours.get(variant, colours["teal"])

    p = doc.add_paragraph()
    set_para_shading(p, bg_hex)
    set_para_border_left(p, border_hex, sz=18, space=8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Mm(2)
    p.paragraph_format.right_indent = Mm(2)

    add_run(p, label.upper(), size=6.5, colour=label_colour, bold=True)
    add_run(p, "\n", size=4)
    add_run(p, text, size=8, colour=C["text_secondary"])


def _evidence_domain_card(doc, ed):
    """Render a long-form evidence domain card."""
    # Card border
    tbl = doc.add_table(rows=1, cols=1)
    set_table_borders(tbl, "E2E8F0", 4)
    cell = tbl.cell(0, 0)
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    # Title + epistemic badge
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, ed["name"], size=9, colour=C["text_primary"], bold=True)
    add_run(p, "    ", size=7)
    ep = ed["epistemic"]
    if ep in EP_BADGES:
        bg, tc = EP_BADGES[ep]
        add_badge_run(p, ep.upper(), bg, tc, size=6)

    # Finding
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    add_run(p2, ed["finding"], size=8, colour=C["text_secondary"])

    # Key tension
    if ed.get("tension"):
        p3 = cell.add_paragraph()
        set_para_shading(p3, HEX["callout_warn"])
        set_para_border_left(p3, "C07A1A", sz=12, space=6)
        p3.paragraph_format.space_after = Pt(6)
        p3.paragraph_format.left_indent = Mm(1)
        add_run(p3, "KEY TENSION\n", size=6, colour=C["amber"], bold=True)
        add_run(p3, ed["tension"], size=7.5, colour=C["text_secondary"])

    # Hypothesis tags
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(4)
    p4.paragraph_format.space_after = Pt(2)
    for tag_text, tag_type in ed["tags"]:
        bg, tc = TAG_COLOURS.get(tag_type, ("F5F5F5", C["text_muted"]))
        add_badge_run(p4, tag_text, bg, tc, size=6)
        add_run(p4, "  ", size=6)

    # Source
    p5 = cell.add_paragraph()
    p5.paragraph_format.space_before = Pt(2)
    p5.paragraph_format.space_after = Pt(0)
    add_run(p5, ed["source"], size=6.5, colour=C["text_muted"], italic=True)


def _hypothesis_card_long(doc, h):
    """Render a full hypothesis card."""
    rgb, hex_c = HYP_COLOURS[h["id"]]

    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)

    # Left colour strip
    strip = tbl.cell(0, 0)
    set_cell_shading(strip, hex_c)
    strip.paragraphs[0].add_run("").font.size = Pt(1)
    tc = strip._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="120" w:type="dxa"/>')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

    # Content
    content = tbl.cell(0, 1)
    set_cell_margins(content, top=80, bottom=80, left=120, right=120)

    # Title + status
    p = content.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"{h['id']}: {h['name']}", size=12, colour=C["text_primary"], bold=True)
    add_run(p, "    ", size=8)
    status_badges = {
        "Priced": ("E8F5E9", C["green"]),
        "Evidence Building": ("FFF3E0", C["amber"]),
        "Active": ("FFEBEE", C["red"]),
        "Slow Burn": ("F5F5F5", C["text_muted"]),
    }
    bg, tc_col = status_badges.get(h["status"], ("F5F5F5", C["text_muted"]))
    add_badge_run(p, h["status"].upper(), bg, tc_col, size=6)

    # Score bar
    p2 = content.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, f"{h['score']}%", size=18, colour=rgb, bold=True, font=MONO)
    dir_arrows = {"Rising": " \u2191 Rising", "Awaiting": " \u2192 Awaiting", "Steady": " \u2192 Steady"}
    dir_col = C["red_light"] if h["direction"] == "Rising" else C["text_muted"]
    add_run(p2, dir_arrows.get(h["direction"], ""), size=8, colour=dir_col, bold=True)

    # Visual score bar
    bar_tbl = doc.add_table(rows=0, cols=0)  # placeholder
    # Actually build inside cell
    p_bar = content.add_paragraph()
    p_bar.paragraph_format.space_after = Pt(6)
    # Use Unicode block characters to simulate the bar
    filled = int(h["score"] / 2)
    empty = 50 - filled
    add_run(p_bar, "\u2588" * filled, size=8, colour=rgb)
    add_run(p_bar, "\u2588" * empty, size=8, colour=RGBColor(0xF3, 0xF5, 0xF7))
    # Remove the placeholder table
    doc._body._body.remove(bar_tbl._tbl)

    # Description
    p3 = content.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    add_run(p3, h.get("full_desc", h["desc"]), size=8.5, colour=C["text_secondary"])

    # Requires
    if h.get("requires"):
        _card_list_heading(content, "REQUIRES")
        for item in h["requires"]:
            _card_list_item(content, item, C["deep_teal"], "\u25A0")

    # Supporting
    if h.get("supporting"):
        _card_list_heading(content, "SUPPORTING EVIDENCE")
        for item in h["supporting"]:
            _card_list_item(content, item, C["green"], "\u25CF")

    # Contradicting
    if h.get("contradicting"):
        _card_list_heading(content, "CONTRADICTING EVIDENCE")
        for item in h["contradicting"]:
            _card_list_item(content, item, C["red"], "\u25CF")


def _card_list_heading(cell, text):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=6.5, colour=C["text_muted"], bold=True)


def _card_list_item(cell, text, dot_colour, dot_char="\u25CF"):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Mm(5)
    p.paragraph_format.first_line_indent = Mm(-4)
    add_run(p, f"{dot_char}  ", size=7, colour=dot_colour)
    add_run(p, text, size=8, colour=C["text_secondary"])


def _tripwire_card_long(doc, tw):
    """Render a long-form tripwire card."""
    tbl = doc.add_table(rows=2, cols=2)
    set_table_borders(tbl, "E2E8F0", 4)

    # Header row
    date_cell = tbl.cell(0, 0)
    name_cell = tbl.cell(0, 1)
    for c in [date_cell, name_cell]:
        set_cell_margins(c, top=40, bottom=40, left=100, right=100)
    add_run(date_cell.paragraphs[0], tw["date"], size=9, colour=C["gold"], bold=True, font=MONO)
    p = name_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, tw["name"].upper(), size=7, colour=C["text_muted"], bold=True)

    # Conditions: 2 cells side by side
    for ci, (direction, condition, consequence) in enumerate(tw["conditions"][:2]):
        cell = tbl.cell(1, ci)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
        set_cell_shading(cell, HEX["sidebar"])

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        cond_prefix = "If " if direction == "positive" else "If "
        cond_colour = C["green"] if direction == "positive" else C["red"]
        add_run(p, f"{cond_prefix}{condition}", size=8, colour=cond_colour, bold=True)

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, consequence, size=7.5, colour=C["text_secondary"])

    # Source
    if tw.get("source"):
        # Add source below the conditions as a merged row... or just add a paragraph after the table
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"Source: {tw['source']}", size=6.5, colour=C["text_muted"], italic=True)


# ═══════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Continuum Trinity Investment Report DOCX Generator")
    parser.add_argument("output", help="Output .docx path")
    parser.add_argument("--format", choices=["short", "long"], required=True, help="Report format")
    args = parser.parse_args()

    doc = Document()

    if args.format == "short":
        build_short_form(doc)
    else:
        build_long_form(doc)

    doc.save(args.output)
    print(f"Generated: {args.output}")


if __name__ == "__main__":
    main()
